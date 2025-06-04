"""Multi-format file reader and editor.

This script provides a command line utility to read and overwrite many common
file formats. It relies on a set of open source libraries; if a dependency is
missing the script will notify the user.

Usage:
    python file_reader.py read <path>
    python file_reader.py write <path> [input_file]
When using ``write`` without ``input_file`` the new contents are read from
``stdin``.
"""

from __future__ import annotations

import argparse
import csv
import json
import sys
from pathlib import Path

try:
    import yaml
except Exception:  # PyYAML may not be installed
    yaml = None

try:
    import pandas as pd
except Exception:  # pandas and friends may not be installed
    pd = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx
except Exception:
    docx = None

try:
    from PIL import Image
except Exception:
    Image = None


def read_text(path: Path) -> str:
    """Read a plain text file."""
    return path.read_text(encoding="utf-8")


def read_json(path: Path) -> str:
    data = json.loads(path.read_text())
    return json.dumps(data, indent=2, ensure_ascii=False)


def read_csv(path: Path) -> str:
    if pd is not None:
        df = pd.read_csv(path)
        return df.to_csv(index=False)
    # fallback to csv module
    with path.open(newline="", encoding="utf-8") as fh:
        reader = csv.reader(fh)
        rows = list(reader)
    return "\n".join(",".join(row) for row in rows)


def read_yaml(path: Path) -> str:
    if yaml is None:
        raise RuntimeError("PyYAML is required to read YAML files")
    data = yaml.safe_load(path.read_text())
    return yaml.safe_dump(data, allow_unicode=True)


def read_excel(path: Path) -> str:
    if pd is None:
        raise RuntimeError("pandas is required to read Excel files")
    df = pd.read_excel(path)
    return df.to_csv(index=False)


def read_pdf(path: Path) -> str:
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 is required to read PDF files")
    reader = PyPDF2.PdfReader(str(path))
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return "\n".join(filter(None, text))


def read_docx(path: Path) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required to read DOCX files")
    document = docx.Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def read_image(path: Path) -> str:
    if Image is None:
        raise RuntimeError("Pillow is required to read image files")
    with Image.open(path) as img:
        return f"Image format: {img.format}, size: {img.size}, mode: {img.mode}"


def write_text(path: Path, content: str) -> None:
    """Overwrite a plain text file."""
    path.write_text(content, encoding="utf-8")


def write_json(path: Path, content: str) -> None:
    data = json.loads(content)
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False))


def write_csv(path: Path, content: str) -> None:
    if pd is not None:
        from io import StringIO
        df = pd.read_csv(StringIO(content))
        df.to_csv(path, index=False)
        return
    rows = [row.split(",") for row in content.splitlines()]
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerows(rows)


def write_yaml(path: Path, content: str) -> None:
    if yaml is None:
        raise RuntimeError("PyYAML is required to write YAML files")
    data = yaml.safe_load(content)
    with path.open("w", encoding="utf-8") as fh:
        yaml.safe_dump(data, fh, allow_unicode=True)


def write_excel(path: Path, content: str) -> None:
    if pd is None:
        raise RuntimeError("pandas is required to write Excel files")
    from io import StringIO
    df = pd.read_csv(StringIO(content))
    df.to_excel(path, index=False)


def write_pdf(path: Path, content: str) -> None:
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 is required to write PDF files")
    try:
        from reportlab.pdfgen import canvas
    except Exception:
        raise RuntimeError("reportlab is required to write PDF files")
    c = canvas.Canvas(str(path))
    textobject = c.beginText(40, 800)
    for line in content.splitlines():
        textobject.textLine(line)
    c.drawText(textobject)
    c.save()


def write_docx(path: Path, content: str) -> None:
    if docx is None:
        raise RuntimeError("python-docx is required to write DOCX files")
    document = docx.Document()
    for line in content.splitlines():
        document.add_paragraph(line)
    document.save(str(path))


def write_image(path: Path, content: str) -> None:
    if Image is None:
        raise RuntimeError("Pillow is required to write image files")
    src = Path(content)
    if not src.exists():
        raise ValueError("For images, content must be path to an image file")
    with Image.open(src) as img:
        img.save(path)


READERS = {
    ".txt": read_text,
    ".json": read_json,
    ".csv": read_csv,
    ".yaml": read_yaml,
    ".yml": read_yaml,
    ".xls": read_excel,
    ".xlsx": read_excel,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".png": read_image,
    ".jpg": read_image,
    ".jpeg": read_image,
}


WRITERS = {
    ".txt": write_text,
    ".json": write_json,
    ".csv": write_csv,
    ".yaml": write_yaml,
    ".yml": write_yaml,
    ".xls": write_excel,
    ".xlsx": write_excel,
    ".pdf": write_pdf,
    ".docx": write_docx,
    ".png": write_image,
    ".jpg": write_image,
    ".jpeg": write_image,
}


def read_file(path: Path) -> str:
    reader = READERS.get(path.suffix.lower())
    if not reader:
        raise ValueError(f"Unsupported file extension: {path.suffix}")
    return reader(path)


def write_file(path: Path, content: str) -> None:
    writer = WRITERS.get(path.suffix.lower())
    if not writer:
        raise ValueError(f"Unsupported file extension: {path.suffix}")
    writer(path, content)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="cmd", required=True)

    read_p = sub.add_parser("read", help="Read a file")
    read_p.add_argument("path", type=Path)

    write_p = sub.add_parser("write", help="Overwrite a file")
    write_p.add_argument("path", type=Path, help="File to modify")
    write_p.add_argument("input", type=Path, nargs="?", help="File containing new contents; stdin if omitted")

    args = parser.parse_args(argv)

    try:
        if args.cmd == "read":
            text = read_file(args.path)
            print(text)
        else:
            if args.input is not None:
                content = args.input.read_text(encoding="utf-8")
            else:
                content = sys.stdin.read()
            write_file(args.path, content)
    except Exception as exc:  # show nice error for missing libs
        parser.error(str(exc))
        return 1

    return 0


if __name__ == "__main__":  # pragma: no cover - entry point
    raise SystemExit(main())

